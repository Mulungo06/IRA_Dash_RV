"""
Aplicação Streamlit - Vigilância de IRAs (Influenza & RSV) — Moçambique
Instituto Nacional de Saúde (INS)

Secção 1 : Processamento de dados (SResult + Demográfico) → Excel limpo
Secção 2 : Geração de relatório Word
Secção 3 : Gráficos e visualizações epidemiológicas
"""

import streamlit as st
import pandas as pd
import numpy as np
from openpyxl import load_workbook
from datetime import datetime, timedelta, date
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT

# ============================================================================
# CONFIGURAÇÃO DA PÁGINA
# ============================================================================

st.set_page_config(
    page_title="Vigilância de IRAs — INS Moçambique",
    page_icon="🦠",
    layout="wide"
)

EMBLEM_PATH  = "Emblem_of_Mozambique.svg.png"
CURRENT_DATE = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

# ============================================================================
# FUNÇÕES AUXILIARES
# ============================================================================

def clean_column_names(df):
    """Normaliza nomes de colunas para snake_case sem acentos."""
    def clean_name(s):
        s = str(s).strip().lower()
        for chars, rep in [('àáâãä','a'),('èéêë','e'),('ìíîï','i'),
                            ('òóôõö','o'),('ùúûü','u'),('ç','c'),('ñ','n')]:
            for c in chars:
                s = s.replace(c, rep)
        s = re.sub(r'[^a-z0-9]+', '_', s)
        return re.sub(r'_+', '_', s).strip('_')
    df.columns = [clean_name(c) for c in df.columns]
    return df

def to_date_str(value):
    """Converte qualquer formato de data para YYYY-MM-DD."""
    if pd.isna(value) or value in ['', 'NA', 'NULL', 'None', 'character(0)']:
        return None
    value_str = str(value).strip()
    if re.match(r'^\d{4}-\d{2}-\d{2}', value_str):
        return value_str[:10]
    try:
        n = float(value_str)
        if n > 1000:
            return (datetime(1899, 12, 30) + timedelta(days=n)).strftime('%Y-%m-%d')
    except Exception:
        pass
    if isinstance(value, (datetime, pd.Timestamp)):
        return value.strftime('%Y-%m-%d')
    return None

def detect_target_sheet(filepath):
    """Detecta a sheet do ano mais recente no ficheiro demográfico."""
    wb     = load_workbook(filepath, read_only=True)
    sheets = wb.sheetnames
    year_sheets = [s for s in sheets if re.match(r'^20[0-9]{2}$', s)]
    if not year_sheets:
        return sheets[-1] if sheets else None
    return str(max(int(y) for y in year_sheets))

# ============================================================================
# TABELA DE UNIDADES SANITÁRIAS (IRAS 1-12 + IDS)
# ============================================================================

UNIDADES_IRAS = {
    # Maputo
    'IRAS1':  ('IRAS1_PED',  'Hospital Central de Maputo (HCM) — Pediatria'),
    'IRAS5':  ('IRAS5_ADU',  'Hospital Central de Maputo (HCM) — Adultos'),
    'IRAS2':  ('IRAS2_PED',  'Hospital Geral de Mavalane (HGM) — Pediatria'),
    'IRAS6':  ('IRAS6_ADU',  'Hospital Geral de Mavalane (HGM) — Adultos'),
    'IRAS3':  ('IRAS3_CSM',  'Centro de Saúde de Mavalane (CSM)'),
    'IRAS4':  ('IRAS4',      'Centro de Saúde de Marracuene'),
    # Beira
    'IRAS7':  ('IRAS7_PED',  'Hospital Central da Beira (HCB) — Pediatria'),
    'IRAS8':  ('IRAS8_ADU',  'Hospital Central da Beira (HCB) — Adultos'),
    # Centro
    'IRAS9':  ('IRAS9',      'Centro de Saúde da Ponta Gêa'),
    'IRAS10': ('IRAS10',     'Hospital Provincial de Tete'),
    # Norte
    'IRAS11': ('IRAS11',     'Centro de Saúde N2'),
    'IRAS12': ('IRAS12',     'Hospital Provincial de Pemba'),
}

def classificar_unidade_sanitaria(codigo, local_colheita=None, idade=None):
    """
    Classifica a unidade sanitária a partir do código IRAS/IDS.
    Retorna (codigo_base, nome_unidade).

    Regra de extracção do site IRAS (partilhada com padronizar_codigo):
    - len(digits)==9  → sempre site de 1 dígito (standard)
    - len(digits)==10 → site de 2 dígitos SE d[:2] in {10,11,12} E year válido
    - Abreviado       → site de 2 dígitos SE total_len>=11 E d[:2] in {10,11,12}
    - Caso contrário  → site de 1 dígito
    """
    # Extrair primeiro código IRAS válido (cell pode ter dois códigos separados por espaço)
    raw = str(codigo).strip().upper()
    parts = re.split(r'[\s\?\|]+', raw)
    codigo = next((p for p in parts if re.match(r'^IRAS\d', re.sub(r'[^A-Z0-9]','',p))),
                  raw)
    codigo = re.sub(r'[^A-Z0-9]', '', codigo)

    # ── IDS ──────────────────────────────────────────────────────────────────
    if codigo.startswith('IDS'):
        num_match = re.search(r'IDS[C]?0?(\d)', codigo)
        if num_match:
            n = int(num_match.group(1))
            if n == 1:
                return ('IDS01_CSZ', 'Centro de Saúde do Zimpeto (CSZ)')
            elif n == 2:
                if local_colheita:
                    loc = str(local_colheita).lower()
                    if 'hg' in loc or ('hospital' in loc and 'mavalane' in loc):
                        idade_anos = extrair_valor_idade(str(idade)) if idade else None
                        if idade_anos is not None and idade_anos < 15:
                            return ('IDS02_HGM_PED', 'Hospital Geral de Mavalane (HGM) — Pediatria (IDS)')
                        return ('IDS02_HGM_ADU', 'Hospital Geral de Mavalane (HGM) — Adultos (IDS)')
                    elif 'cs' in loc or 'centro' in loc:
                        return ('IDS02_CSM', 'Centro de Saúde de Mavalane (CSM) — IDS')
                return ('IDS02_CSM', 'Centro de Saúde de Mavalane (CSM) — IDS')
            elif n == 3:
                return ('IDS03_CSZ', 'Centro de Saúde do Zimpeto (CSZ)')
        return ('IDS_CSZ', 'Centro de Saúde do Zimpeto (CSZ)')

    # ── IRAS ─────────────────────────────────────────────────────────────────
    if codigo.startswith('IRAS'):
        m = re.match(r'^IRAS(\d+)$', codigo)
        if m:
            digits = m.group(1)
            total_len = len(codigo)
            KNOWN_2D = {10, 11, 12}

            def _valid_year(s):
                return (len(s) == 4 and s.startswith('20') and 2015 <= int(s) <= 2035) or \
                       (len(s) == 2 and 15 <= int(s) <= 35)

            site_num = None
            # Standard exact lengths
            if len(digits) == 9:
                site_num = int(digits[0])
            elif len(digits) == 10 and int(digits[:2]) in KNOWN_2D:
                # Only accept 2-digit site if the year fragment is valid
                # seq=digits[2:6], year=digits[6:] → year must be valid 4-digit
                yr = digits[6:]
                if len(yr) == 4 and _valid_year(yr):
                    site_num = int(digits[:2])
                else:
                    site_num = int(digits[0])
            elif total_len >= 11 and int(digits[:2]) in KNOWN_2D:
                site_num = int(digits[:2])
            else:
                site_num = int(digits[0])

            key = f'IRAS{site_num}'
            if key in UNIDADES_IRAS:
                cod_base, nome = UNIDADES_IRAS[key]
                return (cod_base, nome)

    return (codigo, f'Unidade não identificada ({codigo})')

# ============================================================================
# FUNÇÕES AUXILIARES DE RELATÓRIO
# ============================================================================

def extrair_valor_idade(idade_str):
    """Converte string de idade (ex: '5a', '3m', '10d') para anos decimais."""
    if not isinstance(idade_str, str):
        return None
    m = re.match(r'(\d+)([amd])', idade_str.lower())
    if m:
        v, u = int(m.group(1)), m.group(2)
        return v if u == 'a' else (v / 12 if u == 'm' else v / 365)
    return None

def classificar_influenza_subtipos(row):
    """Retorna 'POSITIVO: subtipo(s)' ou 'NEGATIVO' com base em Ct < 40."""
    cols = {"InfA":"A","Apdm":"A(H1pdm)","H1pdm":"A(H1pdm)",
            "H3":"A(H3N2)","H5":"A(H5)","H5a":"A(H5a)","H5b":"A(H5b)",
            "H7":"A(H7)","InfB":"B","Vic":"B(Victoria)","Yam":"B(Yamagata)"}
    found = []
    for col, label in cols.items():
        try:
            if float(row.get(col, 'x')) < 40.0:
                found.append(label)
        except (ValueError, TypeError):
            pass
    return ("POSITIVO: " + ", ".join(found)) if found else "NEGATIVO"

def gerar_resumo_dinamico(df_atual, df_anterior, periodo_atual_str, periodo_anterior_str):
    """
    Gera o texto de resumo epidemiológico por unidade sanitária,
    com desagregação completa por subtipo (Influenza A/B + subtipos, RSV A/B, SARS-CoV-2).
    """
    total = len(df_atual)
    resumo = (f"No período de {periodo_atual_str} foram recebidas e processadas "
              f"{total} {'amostra' if total == 1 else 'amostras'} de casos de "
              f"Infecção Respiratória Aguda (IRA) nos sítios sentinela.\n\n"
              f"Distribuição por unidade sanitária:\n")

    # Classificar unidades
    if 'codigo_unidade' not in df_atual.columns:
        classif = df_atual.apply(
            lambda r: classificar_unidade_sanitaria(
                r.get('Código',''), r.get('Local de Colheita'), r.get('Idade')), axis=1)
        df_atual = df_atual.copy()
        df_atual['codigo_unidade'] = [c[0] for c in classif]
        df_atual['nome_unidade']   = [c[1] for c in classif]

    unidades = df_atual.groupby('codigo_unidade')['nome_unidade'].first().to_dict()
    linhas = []

    def _is_pos(series):
        """Conta positivos numa coluna de resultado."""
        return int(series.astype(str).str.upper().str.contains('POSITIVO', na=False).sum())

    def _tested(series):
        """Conta amostras testadas (não '-')."""
        return int((series.astype(str).str.strip() != '-').sum())

    def _sub_pos(series):
        """Conta positivos numa coluna de subtipo (Positivo/Negativo)."""
        return int(series.astype(str).str.upper().str.strip().eq('POSITIVO').sum())

    for cod_un, nome_un in sorted(unidades.items()):
        df_u = df_atual[df_atual['codigo_unidade'] == cod_un]
        if df_u.empty:
            continue
        n = len(df_u)

        # ── INFLUENZA ────────────────────────────────────────────────────────
        flu_test = _tested(df_u["Influenza"])
        flu_pos  = _is_pos(df_u["Influenza"])

        if flu_test == 0:
            flu_txt = "Influenza: não testada"
        else:
            pct_flu = round(flu_pos / flu_test * 100, 1)
            if flu_pos == 0:
                flu_txt = f"Influenza: 0/{flu_test} positivos (0%)"
            else:
                # Contagem por tipo A/B
                infa_pos  = _sub_pos(df_u["InfA"])  if "InfA"  in df_u.columns else 0
                infb_pos  = _sub_pos(df_u["InfB"])  if "InfB"  in df_u.columns else 0
                h1pdm_pos = _sub_pos(df_u["H1pdm"]) if "H1pdm" in df_u.columns else 0
                h3_pos    = _sub_pos(df_u["H3"])    if "H3"    in df_u.columns else 0
                vic_pos   = _sub_pos(df_u["Vic"])   if "Vic"   in df_u.columns else 0

                # Linha principal
                flu_txt = f"Influenza: {flu_pos}/{flu_test} ({pct_flu}%)"

                # Tipo A
                partes_a = []
                if h1pdm_pos > 0: partes_a.append(f"{h1pdm_pos} H1N1pdm09")
                if h3_pos    > 0: partes_a.append(f"{h3_pos} H3N2")
                if infa_pos  > 0:
                    resto_a = infa_pos - h1pdm_pos - h3_pos
                    if resto_a > 0: partes_a.append(f"{resto_a} não subtipado(s)")
                    flu_txt += f" — Influenza A: {infa_pos}"
                    if partes_a: flu_txt += f" ({', '.join(partes_a)})"

                # Tipo B
                partes_b = []
                if vic_pos > 0: partes_b.append(f"{vic_pos} Victoria")
                if infb_pos > 0:
                    resto_b = infb_pos - vic_pos
                    if resto_b > 0: partes_b.append(f"{resto_b} não subtipado(s)")
                    flu_txt += f"; Influenza B: {infb_pos}"
                    if partes_b: flu_txt += f" ({', '.join(partes_b)})"

        # ── RSV ──────────────────────────────────────────────────────────────
        rsv_test = _tested(df_u["RSV"])
        rsv_pos  = _is_pos(df_u["RSV"])

        if rsv_test == 0:
            rsv_txt = "RSV: não testado"
        else:
            pct_rsv = round(rsv_pos / rsv_test * 100, 1)
            if rsv_pos == 0:
                rsv_txt = f"RSV: 0/{rsv_test} positivos (0%)"
            else:
                rsva_pos = _sub_pos(df_u["RSV_A"]) if "RSV_A" in df_u.columns else 0
                rsvb_pos = _sub_pos(df_u["RSVB"])  if "RSVB"  in df_u.columns else 0
                rsv_txt  = f"RSV: {rsv_pos}/{rsv_test} ({pct_rsv}%)"
                partes_rsv = []
                if rsva_pos > 0: partes_rsv.append(f"{rsva_pos} RSV-A")
                if rsvb_pos > 0: partes_rsv.append(f"{rsvb_pos} RSV-B")
                resto_rsv = rsv_pos - rsva_pos - rsvb_pos
                if resto_rsv > 0: partes_rsv.append(f"{resto_rsv} não subtipado(s)")
                if partes_rsv: rsv_txt += f" ({', '.join(partes_rsv)})"

        # ── SARS-CoV-2 ───────────────────────────────────────────────────────
        sars_test = _tested(df_u["SARS-CoV-2"])
        sars_pos  = _is_pos(df_u["SARS-CoV-2"])
        if sars_test == 0:
            sars_txt = "SARS-CoV-2: não testado"
        else:
            pct_s = round(sars_pos / sars_test * 100, 1)
            sars_txt = f"SARS-CoV-2: {sars_pos}/{sars_test} ({pct_s}%)"

        linhas.append(
            f"  • {nome_un}: {n} amostra{'s' if n!=1 else ''} — "
            f"{flu_txt}; {rsv_txt}; {sars_txt}."
        )

    resumo += "\n".join(linhas) if linhas else "  Nenhuma unidade com dados para o período.\n"

    # ── Comparação semanal ────────────────────────────────────────────────────
    def taxas(df):
        def pct(col):
            t = _tested(df[col]) if col in df.columns else 0
            p = _is_pos(df[col]) if col in df.columns else 0
            return round(100 * p / t, 1) if t > 0 else 0.0
        return pct("Influenza"), pct("SARS-CoV-2"), pct("RSV")

    fa, sa, ra = taxas(df_atual)
    fp, sp, rp = taxas(df_anterior)

    def seta(atual, prev):
        return "↑" if atual > prev else ("↓" if atual < prev else "→")

    resumo += (
        f"\n\nComparação com a semana anterior ({periodo_anterior_str}):\n"
        f"  • Influenza:   {fp}% → {fa}% {seta(fa, fp)}\n"
        f"  • SARS-CoV-2: {sp}% → {sa}% {seta(sa, sp)}\n"
        f"  • RSV:         {rp}% → {ra}% {seta(ra, rp)}"
    )
    return resumo

def criar_tabelas_unidades_sanitarias(doc, df):
    """Gera uma tabela Word por unidade sanitária com resultados individuais."""
    if 'codigo_unidade' not in df.columns:
        classif = df.apply(
            lambda r: classificar_unidade_sanitaria(
                r.get('Código',''), r.get('Local de Colheita'), r.get('Idade')), axis=1)
        df = df.copy()
        df['codigo_unidade'] = [c[0] for c in classif]
        df['nome_unidade']   = [c[1] for c in classif]

    unidades = df.groupby('codigo_unidade')['nome_unidade'].first().to_dict()
    tab_num  = 1
    n_tabs   = 0

    for cod in sorted(unidades):
        df_site = df[df['codigo_unidade'] == cod]
        if df_site.empty:
            continue
        n_tabs += 1
        nome = unidades[cod]

        # Título da tabela
        p = doc.add_paragraph(
            f"Tabela {tab_num}.  Resultados de testagem das amostras provenientes de {nome}",
            style='Heading 2')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        table = doc.add_table(rows=1, cols=10)
        table.style     = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdrs = ["Ord.", "Código", "Sexo", "Idade", "Residência/Bairro",
                "Data de Colheita", "Tipo de Amostra", "Influenza", "RSV", "SARS-CoV-2"]
        for i, h in enumerate(hdrs):
            table.rows[0].cells[i].text = h

        for idx, row in enumerate(df_site.itertuples(index=False), start=1):
            cells = table.add_row().cells
            cells[0].text = str(idx)
            cells[1].text = str(row[0])   # Código
            cells[2].text = str(row[1])   # Sexo
            cells[3].text = str(row[2])   # Idade
            cells[4].text = str(row[3])   # Residência/Bairro
            cells[5].text = row[4].strftime('%d/%m/%Y') if pd.notnull(row[4]) else ""
            cells[6].text = "Swab nasal/orofaríngeo"
            cells[7].text = str(row[7])   # Influenza  (índice pós-remoção Data entrada)
            cells[8].text = str(row[8])   # RSV
            cells[9].text = str(row[9])   # SARS-CoV-2

            # Negrito + vermelho para positivos
            for ci in [7, 8, 9]:
                if "POSITIVO" in cells[ci].text.upper():
                    for para in cells[ci].paragraphs:
                        for run in para.runs:
                            run.font.bold      = True
                            run.font.color.rgb = RGBColor(255, 0, 0)

        # Centralizar
        for r in table.rows:
            for cell in r.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Legenda
        leg = doc.add_paragraph()
        leg.add_run(
            "Legenda — Idade: a = anos, m = meses; RSV: Vírus Sincicial Respiratório; "
            "SARS-CoV-2: coronavírus associado à Síndrome Respiratória Aguda Grave; "
            "\"–\" = amostra não testada para este patógeno."
        ).italic = True
        leg.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in leg.runs:
            run.font.size = Pt(8)

        tab_num += 1

    if n_tabs == 0:
        p = doc.add_paragraph(
            "Nenhuma unidade sanitária com dados para o período seleccionado.",
            style="Heading 2")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return doc

def gerar_relatorio(df_atual, df_anterior, periodo_atual_str, periodo_anterior_str,
                    data_emissao, nome_usuario):
    """Gera o documento Word do relatório analítico semanal de IRAs."""
    doc = Document()

    from docx.enum.section import WD_ORIENT
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = section.bottom_margin = Inches(0.9)
        section.left_margin = section.right_margin = Inches(0.9)

    # Cabeçalho
    header_para = doc.sections[0].header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        header_para.add_run().add_picture(EMBLEM_PATH, width=Inches(0.9))
    except Exception:
        header_para.add_run("[Emblema]")
    header_para.add_run(
        "\nREPÚBLICA DE MOÇAMBIQUE\n"
        "MINISTÉRIO DA SAÚDE\n"
        "INSTITUTO NACIONAL DE SAÚDE (INS)\n"
    )

    # Título
    t = doc.add_heading("RELATÓRIO SEMANAL DE VIGILÂNCIA DE IRAs", level=1)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.add_run(
        "\nVigilância Sentinela de Influenza, RSV e SARS-CoV-2\n"
        "Laboratório Nacional de Referência — INS, Maputo\n"
    ).bold = True

    doc.add_paragraph(
        f"Período do relatório: {periodo_atual_str}",
        style='Heading 2'
    ).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Resumo narrativo
    resumo = gerar_resumo_dinamico(df_atual, df_anterior, periodo_atual_str, periodo_anterior_str)
    doc.add_paragraph(resumo).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph()  # espaço

    # Tabelas por unidade
    df_sem_entrada = df_atual.drop(columns=["Data de entrada"], errors="ignore")
    doc = criar_tabelas_unidades_sanitarias(doc, df_sem_entrada)

    # Rodapé / assinatura
    rodape = doc.add_paragraph()
    rodape.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    rodape.add_run(f"\nData de emissão: {data_emissao}")
    rodape.add_run(f"\nElaborado por: {nome_usuario}")
    rodape.add_run(f"\nGerado automaticamente pelo Sistema de Vigilância de IRAs — INS")
    rodape.add_run(f"\nTimestamp do sistema: {CURRENT_DATE}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ============================================================================
# PROCESSAMENTO INFLUENZA
# ============================================================================

def process_influenza(sresult_file, demo_file, progress_callback=None, target_sheet=None):
    logs = []
    def log(msg):
        logs.append(msg)
        if progress_callback:
            progress_callback(msg)

    log("═══ INICIANDO PROCESSAMENTO — INFLUENZA ═══")

    # ── 1. LER SRESULT ──────────────────────────────────────────────────────
    log("[1/7] A ler SResult...")
    try:
        header_df = pd.read_excel(
            sresult_file, sheet_name="Resultados da Análise",
            header=None, skiprows=2, nrows=2)
        sec_row = header_df.iloc[0].fillna('').astype(str).tolist()
        col_row = header_df.iloc[1].fillna('').astype(str).tolist()
        sec_fill, current = [], ""
        for v in sec_row:
            if v and v.strip():
                current = v.strip()
            sec_fill.append(current)
        log(f"  Cabeçalho: {len(col_row)} colunas detectadas")
    except Exception as e:
        raise Exception(f"Erro ao ler cabeçalho do SResult: {e}")

    try:
        dados_df = pd.read_excel(
            sresult_file, sheet_name="Resultados da Análise",
            header=None, skiprows=4)
        if len(dados_df) == 0:
            raise Exception("SResult sem linhas de dados")
        # Re-ler colunas booleanas como texto
        bool_cols = [i for i, col in enumerate(dados_df.columns) if dados_df[col].dtype == 'bool']
        if bool_cols:
            txt = pd.read_excel(sresult_file, sheet_name="Resultados da Análise",
                                header=None, skiprows=4, usecols=bool_cols, dtype=str)
            for k, ci in enumerate(bool_cols):
                dados_df.iloc[:, ci] = txt.iloc[:, k]
        dados_df.columns = [f"C{i+1}" for i in range(len(dados_df.columns))]
        log(f"  Dados: {len(dados_df)} linhas")
    except Exception as e:
        raise Exception(f"Erro ao ler dados do SResult: {e}")

    # ── 2. LOCALIZAR COLUNAS ────────────────────────────────────────────────
    log("[2/7] A localizar colunas no cabeçalho...")

    def find_col(sec_pat, col_pat, occ=1):
        matches = [i+1 for i,(s,c) in enumerate(zip(sec_fill, col_row))
                   if re.search(sec_pat, s, re.I) and re.search(col_pat, c, re.I)]
        return matches[occ-1] if len(matches) >= occ else None

    i_lab  = next((i+1 for i,c in enumerate(col_row) if c.strip() == 'Laboratório'), None)
    i_ref  = (find_col("Admiss|Presença", r"Número.*Refer|Nº.*Refer|N.*Refer") or
              find_col("NOMES PARA", r"Nº AMOSTRA"))
    i_prov = (find_col("NOMES PARA", r"Província") or
              find_col("NOMES PARA", r"Provincia"))
    i_us   = (find_col("NOMES PARA", r"Unidade de Saúde|Unidade de Saude") or
              find_col("CAPTURA",    r"UNIDADE DE SAUDE"))
    i_col_dt = find_col(r"^Captura$",   r"^dd-mm-aaaa$")
    i_col_hr = find_col(r"Captura",     r"hh:mm")
    i_ddn    = (find_col(r"Admiss|Presença", r"DdN|Data.*Nasc") or
                find_col(r"NOMES PARA",     r"Data de Nascimento"))
    i_idade  = (find_col(r"Admiss|Presença", r"^Idade$") or
                find_col(r"NOMES PARA",     r"^Idade$"))
    i_sexo   = (find_col(r"Admiss|Presença", r"^Sexo$") or
                find_col(r"NOMES PARA",     r"^Sexo$"))
    i_end    = (find_col(r"Admiss|Presença", r"Endereço|Endereco|Residência|Residencia") or
                find_col(r"NOMES PARA",     r"Endereço|Endereco"))
    i_ent_dt = (find_col(r"Captura",         r"Data de entrada") or
                find_col(r"Recibido|Recebido", r"dd-mm-aaaa") or
                find_col(r"Acessada",        r"dd-mm-aaaa"))

    # Gripe / Influenza
    i_flu_test  = find_col(r"RT-PCR INFLU|RTRP-BYAGA", r"^dd-mm-aaaa$", 1)
    i_flu_valid = find_col(r"RT-PCR INFLU|RTRP-BYAGA", r"^dd-mm-aaaa$", 2)
    i_flu_a     = find_col(r"RTRP-FLA",   r"FLU A")
    i_flu_b     = find_col(r"RTRP-FLB",   r"FLU B")
    i_subtipo   = find_col(r"RTRP-SUBTI", r"Subtipo")

    # SARS-CoV-2
    i_sars_test  = find_col(r"RT-PCR SARS|RTRP-SARS|RT-PCR SARS-COV", r"^dd-mm-aaaa$", 1)
    i_sars_valid = find_col(r"RT-PCR SARS|RTRP-SARS|RT-PCR SARS-COV", r"^dd-mm-aaaa$", 2)
    i_sars_res   = (find_col(r"RTCO2|RTRP-SARS", r"^SARS-CoV-2$|^Result$") or
                   find_col(r"RTCO2", r"Result"))

    # RSV
    i_rsv_test  = find_col(r"RT-PCR RSV|RTRP-RSV|RSV-RSVRS|V.RUS SINCIC", r"^dd-mm-aaaa$", 1)
    i_rsv_valid = find_col(r"RT-PCR RSV|RTRP-RSV|RSV-RSVRS|V.RUS SINCIC", r"^dd-mm-aaaa$", 2)
    i_rsv_res   = (find_col(r"RSV-RSVRS", r"RSV Result|RSV") or
                  find_col(r"RTRP-RSV",  r"RSV"))

    log(f"  ref={i_ref}, us={i_us}, idade={i_idade}, sexo={i_sexo}, end={i_end}")
    log(f"  lab={i_lab}, flu_a={i_flu_a}, flu_b={i_flu_b}, rsv={i_rsv_res}, ent_dt={i_ent_dt}")
    if i_ent_dt is None:
        log("  ⚠ 'Data de entrada' não encontrada → usando 'Data da Colheita' como substituto")

    # ── 3. EXTRAIR DADOS ────────────────────────────────────────────────────
    log("[3/7] A montar DataFrame de resultados...")

    def get_col(idx):
        if idx is None or idx < 1 or idx > len(dados_df.columns):
            return pd.Series([None]*len(dados_df), dtype=str)
        col = dados_df[f"C{idx}"]
        if pd.api.types.is_datetime64_any_dtype(col):
            return col.dt.strftime('%Y-%m-%d')
        return col.astype(str)

    # Hora de colheita
    if i_col_hr:
        hr_raw = dados_df[f"C{i_col_hr}"]
        if pd.api.types.is_numeric_dtype(hr_raw):
            hr_fmt = hr_raw.apply(lambda x:
                f"{int(x*24):02d}:{int((x*24-int(x*24))*60):02d}"
                if pd.notna(x) and 0 <= x < 1 else None)
        else:
            hr_fmt = hr_raw.astype(str).str.extract(r'(\d{1,2}:\d{2})')[0]
    else:
        hr_fmt = pd.Series([None]*len(dados_df))

    subtipo  = get_col(i_subtipo)
    rsv_raw  = get_col(i_rsv_res)
    sars_res = get_col(i_sars_res).replace('', None)

    df_sresult = pd.DataFrame({
        'semana_epidemiologica': None,
        'codigo_do_site':        get_col(i_ref),
        'codigo_do_lab':         get_col(i_lab),
        'local_de_colheita':     get_col(i_us),
        'provincia':             get_col(i_prov),
        'data_da_colheita':      get_col(i_col_dt).apply(to_date_str),
        'hora_de_colheita':      hr_fmt,
        'data_de_nascimento':    get_col(i_ddn).apply(to_date_str),
        'idade':                 get_col(i_idade),
        'sexo':                  get_col(i_sexo),
        'residencia_bairro':     get_col(i_end),
        'data_de_entrada':       (get_col(i_ent_dt).apply(to_date_str)
                                  if i_ent_dt else get_col(i_col_dt).apply(to_date_str)),
        'data_da_testagem_flu':  get_col(i_flu_test).apply(to_date_str),
        'data_de_validacao_flu': get_col(i_flu_valid).apply(to_date_str),
        'data_de_testagem_sars': get_col(i_sars_test).apply(to_date_str),
        'data_de_validacao_sars':get_col(i_sars_valid).apply(to_date_str),
        'data_da_testagem_rsv':  get_col(i_rsv_test).apply(to_date_str),
        'data_de_validacao_rsv': get_col(i_rsv_valid).apply(to_date_str),
        'infa':   get_col(i_flu_a),
        'apdm':   None, 'h1pdm': None, 'h3': None,
        'h5': 'N/A', 'h5a': 'N/A', 'h5b': 'N/A', 'h7': 'N/A',
        'infb':   get_col(i_flu_b),
        'vic': None, 'yam': 'N/A',
        'resultado_flu':  None,
        'resultado_sars': sars_res,
        'resultado_rsv':  rsv_raw,
        'rsv_a': None, 'rsvb': None,
        'trl_real_flu': None, 'trl_sars_cov_2': None, 'trl_real_rsv': None,
    })

    # ── 4. COLUNAS DERIVADAS ────────────────────────────────────────────────
    log("[4/7] A calcular subtipos e resultados...")

    df_sresult['apdm']  = subtipo.apply(lambda x: 'Positivo' if 'A(H1pdm)' in str(x) else 'Negativo')
    df_sresult['h1pdm'] = df_sresult['apdm']
    df_sresult['h3']    = subtipo.apply(lambda x: 'Positivo' if 'Sazonal H3N2' in str(x) or 'H3N2' in str(x) else 'Negativo')
    df_sresult['vic']   = subtipo.apply(lambda x: 'Positivo' if 'B-VITORIA' in str(x) or 'VICTORIA' in str(x).upper() else 'Negativo')

    def calc_flu(row):
        if row['infa'] == 'Positivo' or row['infb'] == 'Positivo':
            return 'Positivo'
        if row['infa'] == 'Negativo' or row['infb'] == 'Negativo':
            return 'Negativo'
        return None

    df_sresult['resultado_flu'] = df_sresult.apply(calc_flu, axis=1)

    def std_result(val):
        if pd.isna(val) or str(val).strip() == '':
            return None
        v = str(val).lower()
        if 'positivo' in v: return 'Positivo'
        if 'negativo' in v: return 'Negativo'
        return None

    def std_rsv(val):
        if pd.isna(val) or str(val).strip() == '':
            return None
        v = str(val).lower()
        if 'positivo' in v: return 'Positivo'
        if 'negativo' in v: return 'Negativo'
        if 'infec' in v:    return 'Co-infecção'
        return None

    df_sresult['resultado_sars'] = df_sresult['resultado_sars'].apply(std_result)
    df_sresult['resultado_rsv']  = df_sresult['resultado_rsv'].apply(std_rsv)
    df_sresult['semana_epidemiologica'] = pd.to_datetime(
        df_sresult['data_da_colheita'], errors='coerce').dt.isocalendar().week.astype(str)

    # ── 5. PADRONIZAR CÓDIGO DO SITE ────────────────────────────────────────
    log("[5/7] A padronizar códigos de sítio...")

    _KNOWN_2D = {10, 11, 12}

    def _parse_year(s):
        """Converte sufixo de ano para 4 dígitos. Aceita 2, 3 ou 4 dígitos."""
        s = str(s)
        if len(s) == 4 and s.startswith('20') and 2015 <= int(s) <= 2035:
            return s
        if len(s) == 3 and s[0] == '0' and 2015 <= int('2' + s) <= 2035:
            return '2' + s        # 025 → 2025, 026 → 2026
        if len(s) == 2 and 15 <= int(s) <= 35:
            return '20' + s       # 25 → 2025, 26 → 2026
        return None

    def _try_split(site_str, all_digits):
        """Tenta decompor dígitos após IRAS em (site, seq_padded, year4).
        Aceita seq de 1-4 dígitos e ano de 2-4 dígitos.
        Rejeita seq > 4 dígitos (provável erro de digitação)."""
        rest = all_digits[len(site_str):]
        if len(rest) < 3:
            return None
        for yr_len in [4, 3, 2]:
            if len(rest) > yr_len:
                yr = _parse_year(rest[-yr_len:])
                if yr:
                    seq = rest[:-yr_len]
                    if seq and len(seq) <= 4:   # seq máx 4 dígitos
                        return (int(site_str), seq.zfill(4), yr)
        return None

    def _extrair_primeiro_iras(raw):
        """
        Extrai o primeiro código IRAS válido de uma string que pode conter:
        - Duplicados: 'IRAS702522025 IRAS702522025'
        - Lixo:       'IRAS703222025???IRAS7042025'
        - Espaços/separadores em qualquer posição
        Retorna o primeiro código IRAS{dígitos} encontrado, ou None.
        """
        matches = re.findall(r'IRAS(\d{6,12})', raw.upper())
        if matches:
            return 'IRAS' + matches[0]
        return None

    def padronizar_codigo(codigo):
        """
        Padroniza IRAS para IRAS{site}{seq4}{year4}.

        Formato base: IRAS{1-12}{4 dígitos seq}{4 dígitos ano}
        Variações suportadas:
          - Ano abreviado 2 dígitos: '26' → '2026'
          - Ano abreviado 3 dígitos: '026' → '2026'
          - Seq curta: '29' → '0029'
          - Duplicados ou lixo: extrai o primeiro código IRAS válido
        Seq > 4 dígitos: provável erro de entrada, devolvido sem alterar.
        IDS: nunca alterado.
        """
        if pd.isna(codigo) or str(codigo).strip() in ['', 'None', 'nan']:
            return codigo

        raw = str(codigo).strip()

        # IDS: não alterar nunca
        if raw.upper().startswith('IDS'):
            return re.sub(r'[\s\-\/\.]', '', raw.upper())

        # Se o código contém múltiplos IRAS ou caracteres inválidos, extrair o primeiro
        limpo = re.sub(r'[\s\-\/\.]', '', raw.upper())
        if not re.match(r'^IRAS\d+$', limpo):
            extraido = _extrair_primeiro_iras(raw)
            if extraido:
                limpo = extraido
            else:
                return raw  # não consegue parsear, devolver original

        m = re.match(r'^IRAS(\d+)$', limpo)
        if not m:
            return raw

        all_digits = m.group(1)
        total_len  = len(limpo)

        # Comprimentos standard exactos
        if len(all_digits) == 9:        # site 1d + seq4 + year4
            r = _try_split(all_digits[0], all_digits)
            return f'IRAS{r[0]}{r[1]}{r[2]}' if r else limpo
        if len(all_digits) == 10 and int(all_digits[:2]) in _KNOWN_2D:  # site 2d + seq4 + year4
            r = _try_split(all_digits[:2], all_digits)
            return f'IRAS{r[0]}{r[1]}{r[2]}' if r else limpo

        # Abreviados: site 2d (10-12) só se total_len >= 11
        r2 = None
        if total_len >= 11 and int(all_digits[:2]) in _KNOWN_2D:
            r2 = _try_split(all_digits[:2], all_digits)

        r1 = _try_split(all_digits[0], all_digits)

        if r2:
            return f'IRAS{r2[0]}{r2[1]}{r2[2]}'
        if r1:
            return f'IRAS{r1[0]}{r1[1]}{r1[2]}'
        return limpo  # devolver código limpo mesmo que não parseable

    df_sresult['codigo_do_site'] = df_sresult['codigo_do_site'].apply(padronizar_codigo)

    # ── 6. LER DEMOGRÁFICO ──────────────────────────────────────────────────
    log("[6/7] A ler base demográfica...")
    try:
        if target_sheet is None:
            target_sheet = detect_target_sheet(demo_file)
        log(f"  Sheet seleccionada: {target_sheet}")
        df_demo = pd.read_excel(demo_file, sheet_name=target_sheet, dtype=str)
        df_demo = clean_column_names(df_demo)
        for col in df_demo.columns:
            if 'data' in col:
                df_demo[col] = df_demo[col].apply(to_date_str)
        log(f"  Demográfico: {len(df_demo)} registos, {len(df_demo.columns)} colunas")
    except Exception as e:
        raise Exception(f"Erro ao ler demográfico: {e}")

    # ── 7. COMBINAR ─────────────────────────────────────────────────────────
    log("[7/7] A combinar e filtrar dados...")
    all_cols = set(df_demo.columns) | set(df_sresult.columns)
    for col in all_cols:
        if col not in df_demo.columns:   df_demo[col]    = None
        if col not in df_sresult.columns: df_sresult[col] = None

    combined = pd.concat([df_demo, df_sresult], ignore_index=True)
    n_antes  = len(combined)
    combined = combined[combined['codigo_do_lab'].str.startswith('PMB0', na=False)]
    log(f"  Filtro PMB0: {n_antes} → {len(combined)} registos")

    combined.drop(columns=[c for c in ['trl_sars_co_v_2','inf_a','inf_b'] if c in combined.columns],
                  inplace=True, errors='ignore')

    col_order = [
        'semana_epidemiologica','codigo_do_site','codigo_do_lab',
        'local_de_colheita','provincia','data_da_colheita','hora_de_colheita',
        'data_de_nascimento','idade','sexo','residencia_bairro','data_de_entrada',
        'data_da_testagem_flu','data_de_validacao_flu','data_de_testagem_sars',
        'data_de_validacao_sars','data_da_testagem_rsv','data_de_validacao_rsv',
        'infa','apdm','h1pdm','h3','h5','h5a','h5b','h7',
        'infb','vic','yam','resultado_flu','resultado_sars','resultado_rsv',
        'rsv_a','rsvb','trl_real_flu','trl_sars_cov_2','trl_real_rsv',
    ]
    for col in col_order:
        if col not in combined.columns:
            combined[col] = None
    combined = combined[col_order]

    log(f"═══ CONCLUÍDO: {len(combined)} registos ═══")
    return combined, logs

# ============================================================================
# PROCESSAMENTO RSV
# ============================================================================

def process_rsv(sresult_file, demo_file, progress_callback=None, target_sheet=None):
    logs = []
    def log(msg):
        logs.append(msg)
        if progress_callback:
            progress_callback(msg)

    log("═══ INICIANDO PROCESSAMENTO — RSV ═══")

    log("[1/7] A importar SResult RSV...")
    try:
        df_sr = pd.read_excel(sresult_file, sheet_name="Resultados da Análise", skiprows=3)
        df_sr = clean_column_names(df_sr)
        if len(df_sr) == 0:
            raise Exception("SResult RSV está vazio")
        log(f"  {len(df_sr)} linhas carregadas")
    except Exception as e:
        raise Exception(f"Erro ao ler SResult RSV: {e}")

    log("[2/7] A converter datas...")
    month_map = {'JAN':'01','FEV':'02','MAR':'03','ABR':'04','MAI':'05','JUN':'06',
                 'JUL':'07','AGO':'08','SET':'09','OUT':'10','NOV':'11','DEZ':'12'}
    def cvt_date(d):
        if pd.isna(d) or str(d).strip() == '': return None
        s = str(d).upper()
        for k,v in month_map.items(): s = s.replace(k, v)
        try: return pd.to_datetime(s, format='%d-%m-%Y').strftime('%Y-%m-%d')
        except:
            try: return pd.to_datetime(s).strftime('%Y-%m-%d')
            except: return None
    for col in df_sr.columns:
        if 'data' in col:
            df_sr[col] = df_sr[col].apply(cvt_date)

    log("[3/7] A seleccionar colunas RSV...")
    cols_keep = ['numero_s_de_referencia','laboratorio','unidade_de_saude','provincia_2',
                 'data','idade','sexo','data_2','rsv_result','subtipo']
    cols_found = [c for c in cols_keep if c in df_sr.columns] or list(df_sr.columns)
    df_sr = df_sr[cols_found]

    log("[4/7] A padronizar nomes...")
    df_sr = df_sr.rename(columns={
        'numero_s_de_referencia': 'codigo_do_site',
        'laboratorio':            'codigo_do_lab',
        'unidade_de_saude':       'local_de_colheita',
        'provincia_2':            'provincia',
        'data':                   'data_da_colheita',
        'data_2':                 'data_de_entrada',
        'rsv_result':             'resultado_rsv',
    })
    if 'data_de_entrada' not in df_sr.columns and 'data_da_colheita' in df_sr.columns:
        df_sr['data_de_entrada'] = df_sr['data_da_colheita']
        log("  ⚠ 'data_de_entrada' ausente → usando 'data_da_colheita'")

    log("[5/7] A padronizar resultados RSV...")
    def std_rsv(val):
        if pd.isna(val): return None
        v = str(val).strip().upper()
        if 'POSITIVO' in v: return 'Positivo'
        if 'NEGATIVO' in v: return 'Negativo'
        if 'INFEC' in v or 'CO-INFEC' in v: return 'Co-infecção'
        return v if v else None
    if 'resultado_rsv' in df_sr.columns:
        df_sr['resultado_rsv'] = df_sr['resultado_rsv'].apply(std_rsv)

    log("[6/7] A ler base demográfica RSV...")
    try:
        if target_sheet is None:
            target_sheet = detect_target_sheet(demo_file)
        log(f"  Sheet: {target_sheet}")
        df_demo = pd.read_excel(demo_file, sheet_name=target_sheet, dtype=str)
        df_demo = clean_column_names(df_demo)
        for col in df_demo.columns:
            if 'data' in col:
                df_demo[col] = df_demo[col].apply(to_date_str)
        log(f"  Demográfico: {len(df_demo)} registos")
    except Exception as e:
        raise Exception(f"Erro ao ler demográfico RSV: {e}")

    log("[7/7] A combinar dados RSV...")
    all_cols = set(df_demo.columns) | set(df_sr.columns)
    for col in all_cols:
        if col not in df_demo.columns: df_demo[col] = None
        if col not in df_sr.columns:   df_sr[col]   = None
    combined = pd.concat([df_demo, df_sr], ignore_index=True)
    if 'codigo_do_lab' in combined.columns:
        combined = combined[combined['codigo_do_lab'].str.startswith('PMB0', na=False)]

    log(f"═══ RSV CONCLUÍDO: {len(combined)} registos ═══")
    return combined, logs

# ============================================================================
# SESSION STATE
# ============================================================================

if 'dados_processados' not in st.session_state:
    st.session_state['dados_processados'] = None
if 'tipo_analise' not in st.session_state:
    st.session_state['tipo_analise'] = 'Influenza'
if 'logs' not in st.session_state:
    st.session_state['logs'] = []

# ============================================================================
# SIDEBAR
# ============================================================================

st.sidebar.title("🦠 IRAs — INS Moçambique")
secao = st.sidebar.radio(
    "Navegação",
    ["📊 Processamento de Dados", "📝 Geração de Relatório", "📈 Gráficos e Visualizações"],
    index=0
)
st.sidebar.markdown("---")
st.sidebar.markdown("""
**Fluxo recomendado:**
1. 📊 Carregar SResult + Demográfico e processar
2. 📝 Seleccionar período e gerar relatório Word
3. 📈 Explorar gráficos epidemiológicos
""")
st.sidebar.markdown("---")
st.sidebar.caption(f"v3.0 · INS · {datetime.now().strftime('%Y-%m-%d')}")

# ============================================================================
# SECÇÃO 1 — PROCESSAMENTO DE DADOS
# ============================================================================

if secao == "📊 Processamento de Dados":
    st.title("📊 Processamento de Dados SResult + Demográficos")
    st.markdown(
        "Carregue o ficheiro **SResult** exportado do LIMS e a **base demográfica** "
        "acumulada. O sistema combina os dois, padroniza os campos e gera um Excel "
        "pronto para relatório e análise.")
    st.markdown("---")

    tipo_analise = st.radio(
        "**Patógeno / tipo de análise:**",
        ["Influenza", "RSV"], horizontal=True, key='tipo_radio')
    st.session_state['tipo_analise'] = tipo_analise

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("📄 Ficheiro SResult")
        st.caption(f"Formato esperado: SResult_{tipo_analise}_semanaXX_AAAA.xlsx")
        sresult_file = st.file_uploader(
            "Carregar SResult", type=['xlsx'], key='sresult_upload',
            help="Exportação directa do LIMS — folha 'Resultados da Análise'")
    with col2:
        st.subheader("📄 Base Demográfica")
        st.caption(f"Formato esperado: Dados_demograficos_{tipo_analise}_AAAA.xlsx")
        demo_file = st.file_uploader(
            "Carregar Demográfico", type=['xlsx'], key='demo_upload',
            help="Ficheiro acumulado com histórico de todas as semanas do ano")

    # Selecção de ano/sheet no demográfico
    selected_sheet = None
    if demo_file:
        try:
            wb = load_workbook(demo_file, read_only=True)
            sheets_disp = wb.sheetnames
            wb.close()
            demo_file.seek(0)
            year_sheets = [s for s in sheets_disp if re.match(r'^20[0-9]{2}$', s)]
            if len(year_sheets) > 1:
                selected_sheet = st.selectbox(
                    "📅 Ano a processar (sheet da base demográfica)",
                    options=year_sheets, index=len(year_sheets)-1, key='sheet_select_proc',
                    help="Seleccione o ano correspondente ao período do SResult")
            elif len(year_sheets) == 1:
                selected_sheet = year_sheets[0]
                st.info(f"📅 Sheet detectada automaticamente: **{selected_sheet}**")
            else:
                selected_sheet = sheets_disp[-1] if sheets_disp else None
                st.info(f"📅 Sheet utilizada: **{selected_sheet}**")
        except Exception:
            selected_sheet = None

    st.markdown("---")

    if st.button("🔄 Processar Dados", type="primary", use_container_width=True):
        if not sresult_file or not demo_file:
            st.error("❌ Carregue ambos os ficheiros antes de processar.")
        else:
            progress_bar = st.progress(0, text="A iniciar...")
            status_text  = st.empty()
            log_expander = st.expander("📋 Registo de processamento", expanded=False)

            def update_progress(msg):
                status_text.text(msg)
                st.session_state['logs'].append(msg)

            try:
                if tipo_analise == "Influenza":
                    progress_bar.progress(15, text="A processar Influenza...")
                    df_result, logs = process_influenza(sresult_file, demo_file, update_progress, selected_sheet)
                else:
                    progress_bar.progress(15, text="A processar RSV...")
                    df_result, logs = process_rsv(sresult_file, demo_file, update_progress, selected_sheet)

                progress_bar.progress(100, text="✅ Concluído!")
                st.session_state['dados_processados'] = df_result
                st.session_state['logs']              = logs

                with log_expander:
                    for m in logs:
                        st.text(m)

                st.success(f"✅ Processamento concluído — {len(df_result):,} registos combinados.")

                # Métricas de resumo
                def n_pos(col):
                    return df_result[col].astype(str).str.upper().str.contains('POSITIVO', na=False).sum() \
                        if col in df_result.columns else 0

                pos_flu  = n_pos('resultado_flu')
                pos_sars = n_pos('resultado_sars')
                pos_rsv  = n_pos('resultado_rsv')

                c1,c2,c3,c4,c5 = st.columns(5)
                c1.metric("Total de registos", f"{len(df_result):,}")
                c2.metric("Influenza +",        f"{pos_flu:,}")
                c3.metric("SARS-CoV-2 +",       f"{pos_sars:,}")
                c4.metric("RSV +",               f"{pos_rsv:,}")
                c5.metric("Total positivos",     f"{pos_flu+pos_sars+pos_rsv:,}")

                st.subheader("👁️ Pré-visualização dos dados processados")
                st.dataframe(df_result.head(50), use_container_width=True, height=380)

                st.markdown("---")
                st.subheader("💾 Exportar para Excel")
                buf = BytesIO()
                with pd.ExcelWriter(buf, engine='openpyxl') as w:
                    df_result.to_excel(w, index=False, sheet_name='Dados')
                buf.seek(0)
                st.download_button(
                    label=f"📥 Download Excel — {len(df_result):,} registos",
                    data=buf,
                    file_name=f"dados_{tipo_analise}_{datetime.now().strftime('%Y-%m-%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True, type="primary")
                st.info("💡 Passe à secção **Geração de Relatório** para criar o documento Word, "
                        "ou à secção **Gráficos** para análise visual.")

            except Exception as e:
                progress_bar.empty()
                status_text.empty()
                st.error(f"❌ Erro: {e}")
                with st.expander("🔍 Detalhes técnicos do erro"):
                    import traceback
                    st.code(traceback.format_exc())

    elif st.session_state['dados_processados'] is not None:
        df_mem = st.session_state['dados_processados']
        st.info(f"ℹ️ Já existem dados de **{st.session_state['tipo_analise']}** em memória "
                f"({len(df_mem):,} registos). Pode avançar para as secções seguintes.")
        if st.button("🗑️ Limpar dados e recomeçar"):
            st.session_state['dados_processados'] = None
            st.session_state['logs'] = []
            st.rerun()

# ============================================================================
# SECÇÃO 2 — GERAÇÃO DE RELATÓRIO
# ============================================================================

elif secao == "📝 Geração de Relatório":
    st.title("📝 Geração de Relatório Semanal — Word")
    st.markdown(
        "Gera o relatório analítico semanal em formato Word (.docx) pronto a partilhar, "
        "com cabeçalho institucional, resumo narrativo por unidade sanitária e tabelas "
        "detalhadas de resultados individuais.")
    st.markdown("---")

    st.info("💡 Pode usar os dados processados na secção anterior **ou** carregar um Excel já processado.")

    use_processed = st.checkbox(
        "✅ Usar dados já processados (secção anterior)",
        value=st.session_state['dados_processados'] is not None)

    df_relatorio = None

    if use_processed:
        if st.session_state['dados_processados'] is not None:
            df_relatorio = st.session_state['dados_processados'].copy()
            tipo_rel     = st.session_state['tipo_analise']
            st.success(f"✅ Dados em memória carregados: {len(df_relatorio):,} registos ({tipo_rel})")
        else:
            st.warning("⚠️ Nenhum dado processado em memória. Processe primeiro ou carregue um ficheiro abaixo.")
    else:
        st.subheader("📤 Carregar ficheiro Excel processado")
        uploaded_excel = st.file_uploader(
            "Ficheiro Excel (output do processamento)",
            type=['xlsx'], key='relatorio_upload',
            help="Excel exportado na secção de Processamento de Dados")
        if uploaded_excel:
            try:
                wb_r = load_workbook(uploaded_excel, read_only=True)
                sheets_r = wb_r.sheetnames
                wb_r.close()
                uploaded_excel.seek(0)
                yr = [s for s in sheets_r if re.match(r'^20[0-9]{2}$', s)]
                if len(yr) > 1:
                    sel_r = st.selectbox("📅 Seleccionar ano (sheet)", yr,
                                         index=len(yr)-1, key='sheet_sel_rel')
                    df_relatorio = pd.read_excel(uploaded_excel, sheet_name=sel_r, dtype=str)
                elif yr:
                    df_relatorio = pd.read_excel(uploaded_excel, sheet_name=yr[0], dtype=str)
                else:
                    df_relatorio = pd.read_excel(uploaded_excel)
                tipo_rel = 'Influenza' if ('infa' in df_relatorio.columns or
                                            'infb' in df_relatorio.columns) else 'RSV'
                st.success(f"✅ Ficheiro carregado: {len(df_relatorio):,} registos ({tipo_rel})")
            except Exception as e:
                st.error(f"❌ Erro ao ler ficheiro: {e}")

    if df_relatorio is not None:
        st.markdown("---")
        st.subheader("⚙️ Configurações do relatório")

        def enc(df, *nomes):
            for n in nomes:
                if n in df.columns:
                    return n
            return None

        col_codigo      = enc(df_relatorio, 'codigo_do_site',   'Código do Site')
        col_sexo        = enc(df_relatorio, 'sexo',             'Sexo')
        col_idade       = enc(df_relatorio, 'idade',            'Idade')
        col_residencia  = enc(df_relatorio, 'residencia_bairro','Residência/Bairro')
        col_data_col    = enc(df_relatorio, 'data_da_colheita', 'Data da Colheita')
        col_data_ent    = enc(df_relatorio, 'data_de_entrada',  'Data de entrada')
        col_local       = enc(df_relatorio, 'local_de_colheita','Local de Colheita')
        col_flu         = enc(df_relatorio, 'resultado_flu',    'Resultado Flu')
        col_rsv         = enc(df_relatorio, 'resultado_rsv',    'Resultado RSV ', 'Resultado RSV')
        col_sars        = enc(df_relatorio, 'resultado_sars',   'Resultado  SARS','Resultado SARS')
        col_infa        = enc(df_relatorio, 'infa',   'InfA')
        col_infb        = enc(df_relatorio, 'infb',   'InfB')
        col_h1pdm       = enc(df_relatorio, 'h1pdm',  'apdm', 'Apdm', 'H1pdm')
        col_h3          = enc(df_relatorio, 'h3',     'H3')
        col_vic         = enc(df_relatorio, 'vic',    'Vic')
        col_rsva        = enc(df_relatorio, 'rsv_a',  'RSV A', 'RSV_A')
        col_rsvb        = enc(df_relatorio, 'rsvb',   'RSVB',  'RSV B')

        def fmt_result(col):
            if col is None:
                return pd.Series(["-"]*len(df_relatorio))
            return df_relatorio[col].apply(
                lambda x: str(x).strip().upper() if pd.notna(x) and str(x).strip() not in ['','nan','None'] else "-")

        def _col_str(col):
            """Devolve a coluna como string, ou série de '' se não existir."""
            return df_relatorio[col].astype(str).str.upper().str.strip() if col else pd.Series([""]*len(df_relatorio))

        df_fmt = pd.DataFrame({
            "Código":            df_relatorio[col_codigo].astype(str).str.strip()      if col_codigo   else "",
            "Sexo":              df_relatorio[col_sexo].astype(str).str.upper()         if col_sexo     else "",
            "Idade":             df_relatorio[col_idade].astype(str)                    if col_idade    else "",
            "Residência/Bairro": df_relatorio[col_residencia].astype(str).fillna("N/D") if col_residencia else "N/D",
            "Data da Colheita":  pd.to_datetime(df_relatorio[col_data_col],  errors='coerce') if col_data_col else pd.NaT,
            "Data de entrada":   pd.to_datetime(df_relatorio[col_data_ent],  errors='coerce') if col_data_ent else pd.NaT,
            "Local de Colheita": df_relatorio[col_local].astype(str)                    if col_local    else "",
            "Tipo de Amostra":   "Swab nasal/orofaríngeo",
            "Influenza":         fmt_result(col_flu),
            "RSV":               fmt_result(col_rsv),
            "SARS-CoV-2":        fmt_result(col_sars),
            "InfA":              _col_str(col_infa),
            "InfB":              _col_str(col_infb),
            "H1pdm":             _col_str(col_h1pdm),
            "H3":                _col_str(col_h3),
            "Vic":               _col_str(col_vic),
            "RSV_A":             _col_str(col_rsva),
            "RSVB":              _col_str(col_rsvb),
        })

        # Selecção de período
        datas_val = df_fmt["Data de entrada"].dropna()
        if len(datas_val) > 0:
            min_d = datas_val.min().date()
            max_d = min(datas_val.max().date(), date.today())
        else:
            min_d = date.today() - timedelta(days=30)
            max_d = date.today()

        c1, c2 = st.columns(2)
        with c1:
            data_inicio = st.date_input(
                "Data inicial do período actual",
                value=max_d - timedelta(days=6),
                min_value=date(2020,1,1), max_value=max_d,
                help="Primeira data da semana epidemiológica a reportar")
        with c2:
            data_fim = st.date_input(
                "Data final do período actual",
                value=max_d,
                min_value=data_inicio, max_value=max_d,
                help="Última data da semana epidemiológica a reportar")

        c1, c2 = st.columns(2)
        with c1:
            nome_usuario = st.text_input(
                "Nome do responsável pela emissão", value="",
                placeholder="Ex.: Dr. João Silva — Epidemiologia INS")
        with c2:
            data_emissao = st.text_input(
                "Data de emissão", value=datetime.now().strftime("%d/%m/%Y"))

        # Filtrar período actual e anterior
        df_atual = df_fmt[
            (df_fmt["Data de entrada"] >= pd.Timestamp(data_inicio)) &
            (df_fmt["Data de entrada"] <= pd.Timestamp(data_fim))
        ].copy()

        data_ini_prev = data_inicio - timedelta(days=7)
        data_fim_prev = data_fim    - timedelta(days=7)
        df_anterior   = df_fmt[
            (df_fmt["Data de entrada"] >= pd.Timestamp(data_ini_prev)) &
            (df_fmt["Data de entrada"] <= pd.Timestamp(data_fim_prev))
        ].copy()

        periodo_actual_str   = f"{data_inicio.strftime('%d/%m/%Y')} a {data_fim.strftime('%d/%m/%Y')}"
        periodo_anterior_str = f"{data_ini_prev.strftime('%d/%m/%Y')} a {data_fim_prev.strftime('%d/%m/%Y')}"

        st.info(f"📊 Período actual: **{len(df_atual):,}** registos ({periodo_actual_str})")
        st.info(f"📊 Semana anterior (comparação): **{len(df_anterior):,}** registos ({periodo_anterior_str})")
        st.markdown("---")

        if st.button("📄 Gerar Relatório Word", type="primary", use_container_width=True):
            if len(df_atual) == 0:
                st.error("❌ Nenhum registo no período seleccionado. Ajuste as datas.")
            else:
                with st.spinner("A gerar documento Word..."):
                    try:
                        doc_io = gerar_relatorio(
                            df_atual, df_anterior,
                            periodo_actual_str, periodo_anterior_str,
                            data_emissao, nome_usuario)
                        st.success("✅ Relatório gerado com sucesso!")
                        fname = f"Relatorio_IRAs_{data_inicio.strftime('%Y%m%d')}_{data_fim.strftime('%Y%m%d')}.docx"
                        st.download_button(
                            "📥 Download Relatório Word", data=doc_io.getvalue(),
                            file_name=fname,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True, type="primary")
                        st.markdown("""
**O relatório inclui:**
- ✅ Cabeçalho institucional INS (Emblema + título)
- ✅ Resumo narrativo com positividade por unidade sanitária
- ✅ Comparação com a semana anterior (↑↓→)
- ✅ Tabelas individuais por sítio sentinela (IRAS1–12, IDS)
- ✅ Positivos destacados em **negrito vermelho**
- ✅ Orientação paisagem (landscape) com margens optimizadas
                        """)
                    except Exception as e:
                        st.error(f"❌ Erro ao gerar relatório: {e}")
                        with st.expander("🔍 Detalhes técnicos"):
                            import traceback
                            st.code(traceback.format_exc())

        if st.checkbox("👁️ Pré-visualizar dados do período actual", value=False):
            st.dataframe(df_atual, use_container_width=True)

# ============================================================================
# SECÇÃO 3 — GRÁFICOS E VISUALIZAÇÕES
# ============================================================================

elif secao == "📈 Gráficos e Visualizações":
    st.title("📈 Gráficos e Visualizações Epidemiológicas — IRAs")
    st.markdown(
        "Explore a distribuição temporal, geográfica e demográfica dos casos de IRA "
        "testados na rede sentinela.")
    st.markdown("---")

    st.info("💡 Use os dados processados na secção anterior **ou** carregue um ficheiro Excel.")

    use_pg = st.checkbox(
        "✅ Usar dados processados",
        value=st.session_state['dados_processados'] is not None,
        key='use_processed_graphs')

    df_graphs = None
    if use_pg:
        if st.session_state['dados_processados'] is not None:
            df_graphs = st.session_state['dados_processados'].copy()
            st.success(f"✅ Dados carregados: **{len(df_graphs):,}** registos")
        else:
            st.warning("⚠️ Nenhum dado processado. Processe primeiro ou carregue abaixo.")
    else:
        up_g = st.file_uploader("Carregar ficheiro Excel", type=['xlsx'], key='graphs_upload')
        if up_g:
            try:
                wb_g = load_workbook(up_g, read_only=True)
                sh_g = wb_g.sheetnames
                wb_g.close()
                up_g.seek(0)
                yr_g = [s for s in sh_g if re.match(r'^20[0-9]{2}$', s)]
                if len(yr_g) > 1:
                    sel_g = st.selectbox("📅 Ano (sheet)", yr_g, index=len(yr_g)-1, key='sheet_g')
                    df_graphs = pd.read_excel(up_g, sheet_name=sel_g, dtype=str)
                elif yr_g:
                    df_graphs = pd.read_excel(up_g, sheet_name=yr_g[0], dtype=str)
                else:
                    df_graphs = pd.read_excel(up_g)
                st.success(f"✅ Ficheiro carregado: **{len(df_graphs):,}** registos")
            except Exception as e:
                st.error(f"❌ Erro: {e}")

    if df_graphs is not None:
        # ── Mapeamento de colunas ──────────────────────────────────────────
        COLS_MAP = {
            'codigo_do_site':        ['codigo_do_site','Código do Site','codigo do site'],
            'local_de_colheita':     ['local_de_colheita','Local de Colheita'],
            'provincia':             ['provincia','Província','Provincia'],
            'data_da_colheita':      ['data_da_colheita','Data da Colheita'],
            'data_de_entrada':       ['data_de_entrada','Data de entrada'],
            'semana_epidemiologica': ['semana_epidemiologica','Semana Epidemiologica','Semana Epidemiológica'],
            'idade':                 ['idade','Idade'],
            'sexo':                  ['sexo','Sexo'],
            'infa':                  ['infa','InfA'],
            'infb':                  ['infb','InfB'],
            'apdm':                  ['apdm','Apdm'],
            'h1pdm':                 ['h1pdm','H1pdm'],
            'h3':                    ['h3','H3'],
            'vic':                   ['vic','Vic'],
            'yam':                   ['yam','Yam'],
            'resultado_flu':         ['resultado_flu','Resultado Flu','Resultado FLU'],
            'resultado_sars':        ['resultado_sars','Resultado  SARS','Resultado SARS'],
            'resultado_rsv':         ['resultado_rsv','Resultado RSV','Resultado RSV '],
            'rsv_a':                 ['rsv_a','RSV A','RSV_A'],
            'rsvb':                  ['rsvb','RSVB','RSV B'],
        }
        ren = {}
        for can, alts in COLS_MAP.items():
            if can not in df_graphs.columns:
                for alt in alts:
                    if alt in df_graphs.columns and alt != can:
                        ren[alt] = can
                        break
        if ren:
            df_graphs = df_graphs.rename(columns=ren)

        for col in ['resultado_flu','resultado_sars','resultado_rsv','infa','infb']:
            if col in df_graphs.columns and df_graphs[col].dtype == 'object':
                df_graphs[col] = df_graphs[col].astype(str).str.strip()

        # Usar data_de_entrada como eixo temporal (igual ao relatório)
        col_dt = 'data_de_entrada' if 'data_de_entrada' in df_graphs.columns else 'data_da_colheita'
        if col_dt in df_graphs.columns:
            df_graphs['_dt'] = pd.to_datetime(df_graphs[col_dt], errors='coerce')
        else:
            df_graphs['_dt'] = pd.NaT

        if 'data_da_colheita' in df_graphs.columns:
            df_graphs['data_colheita_dt'] = pd.to_datetime(df_graphs['data_da_colheita'], errors='coerce')

        st.info(f"📊 Total de registos carregados: **{len(df_graphs):,}**")
        st.markdown("---")

        # ── Filtros ───────────────────────────────────────────────────────
        st.subheader("🔍 Filtros")
        fc1, fc2, fc3 = st.columns(3)

        with fc1:
            dv = df_graphs['_dt'].dropna()
            if len(dv) > 0:
                d_min = dv.min().date()
                d_max = dv.max().date()
                d_ini = st.date_input("Data início", value=d_min,
                                      min_value=d_min, max_value=d_max, key='g_di')
                d_fim = st.date_input("Data fim",    value=d_max,
                                      min_value=d_ini, max_value=d_max, key='g_df')
                df_graphs = df_graphs[
                    (df_graphs['_dt'] >= pd.Timestamp(d_ini)) &
                    (df_graphs['_dt'] <= pd.Timestamp(d_fim))]

        with fc2:
            if 'provincia' in df_graphs.columns:
                provs = ['Todas'] + sorted(df_graphs['provincia'].dropna().unique().tolist())
                prov_sel = st.selectbox("Província", provs, key='g_prov')
                if prov_sel != 'Todas':
                    df_graphs = df_graphs[df_graphs['provincia'] == prov_sel]

        with fc3:
            if 'codigo_do_site' in df_graphs.columns:
                un_map = {
                    "Todas":  "Todas as unidades",
                    "IRAS1":  "HCM — Pediatria",
                    "IRAS2":  "HGM — Pediatria",
                    "IRAS3":  "CS Mavalane",
                    "IRAS4":  "CS Marracuene",
                    "IRAS5":  "HCM — Adultos",
                    "IRAS6":  "HGM — Adultos",
                    "IRAS7":  "HCB — Pediatria (Beira)",
                    "IRAS8":  "HCB — Adultos (Beira)",
                    "IRAS9":  "CS Ponta Gêa (Beira)",
                    "IRAS10": "HP Tete",
                    "IRAS11": "CS N2",
                    "IRAS12": "HP Pemba",
                    "IDS":    "CSZ / HGM / CSM (IDS)",
                }
                un_sel = st.selectbox(
                    "Unidade sanitária", list(un_map.keys()),
                    format_func=lambda x: un_map[x], key='g_un')
                if un_sel != 'Todas':
                    df_graphs = df_graphs[df_graphs['codigo_do_site'].str.startswith(un_sel, na=False)]

        st.info(f"📊 Registos após filtros: **{len(df_graphs):,}**")

        # ── Semana epidemiológica — calculada APÓS os filtros, usando data_de_entrada ──
        # Usar _dt (= data_de_entrada) como referência: garante que as semanas
        # correspondem exactamente ao período seleccionado nos filtros.
        # data_da_colheita pode ser anterior ao período (ex: amostras de Dezembro
        # recebidas em Janeiro), o que causaria semanas 49-52 num filtro de Janeiro.
        if '_dt' in df_graphs.columns:
            sem_calc = df_graphs['_dt'].dt.isocalendar().week.astype('Int64').astype(str)
            sem_calc = sem_calc.replace({'<NA>':'', 'nan':'', 'None':''})
            if 'semana_epidemiologica' in df_graphs.columns:
                sem_existente = df_graphs['semana_epidemiologica'].fillna('').astype(str).str.strip()
                # Recalcular sempre a partir de _dt para garantir consistência com o filtro
                df_graphs['semana_epidemiologica'] = df_graphs.apply(
                    lambda r: sem_calc[r.name] if sem_calc[r.name] not in ['','<NA>']
                    else sem_existente[r.name], axis=1)
            else:
                df_graphs['semana_epidemiologica'] = sem_calc
            df_graphs['semana_epidemiologica'] = df_graphs['semana_epidemiologica'].replace(
                ['<NA>','nan','None',''], pd.NA)

        # ── Gráfico 1: Visão geral ────────────────────────────────────────
        st.markdown("---")
        st.header("1️⃣ Visão Geral — Amostras e Positividade")

        tot = len(df_graphs)

        def _is_pos(col):
            if col not in df_graphs.columns:
                return pd.Series([False] * len(df_graphs), index=df_graphs.index)
            return df_graphs[col].astype(str).str.upper().str.contains('POSITIVO', na=False)

        pos_flu  = int(_is_pos('resultado_flu').sum())
        pos_sars = int(_is_pos('resultado_sars').sum())
        pos_rsv  = int(_is_pos('resultado_rsv').sum())

        df_graphs['_qualquer_pos'] = _is_pos('resultado_flu') | _is_pos('resultado_sars') | _is_pos('resultado_rsv')
        tot_pos = int(df_graphs['_qualquer_pos'].sum())
        tot_neg = tot - tot_pos

        g1, g2 = st.columns(2)
        with g1:
            st.subheader("Amostras positivas vs negativas")
            st.plotly_chart({
                'data': [{'labels':['Positivas','Negativas'],'values':[tot_pos,tot_neg],
                          'type':'pie','marker':{'colors':['#ef4444','#22c55e']},
                          'hole':0.42,'textinfo':'label+percent+value',
                          'textposition':'outside'}],
                'layout': {'title':f'Total: {tot:,} amostras','height':380,'showlegend':True}
            }, use_container_width=True, key='gc_pie')

        with g2:
            st.subheader("Positivos por patógeno")
            st.plotly_chart({
                'data': [{'x':['Influenza','SARS-CoV-2','RSV'],'y':[pos_flu,pos_sars,pos_rsv],
                          'type':'bar','marker':{'color':['#3b82f6','#ef4444','#f59e0b']},
                          'text':[pos_flu,pos_sars,pos_rsv],'textposition':'outside'}],
                'layout': {'title':'Positivos por patógeno','height':380,
                           'xaxis':{'title':'Patógeno'},'yaxis':{'title':'N.º de casos'}}
            }, use_container_width=True, key='gc_bar')

        m1,m2,m3,m4 = st.columns(4)
        m1.metric("Total amostras",  f"{tot:,}")
        m2.metric("Total positivas", f"{tot_pos:,}", f"{(tot_pos/tot*100):.1f}%" if tot > 0 else "0%")
        m3.metric("Influenza +",     f"{pos_flu:,}")
        m4.metric("RSV +",           f"{pos_rsv:,}")

        soma = pos_flu + pos_sars + pos_rsv
        if soma > tot_pos:
            st.info(f"ℹ️ Existem **{soma-tot_pos}** co-infecção(ões) (amostra positiva para mais do que um patógeno).")

        # ── Gráfico 2: Por unidade sanitária ─────────────────────────────
        st.markdown("---")
        st.subheader("📍 Amostras por Unidade Sanitária")

        def classif_dir(cod, loc=None, idade=None):
            if pd.isna(cod): return None
            c = str(cod).strip().upper()
            if not c or c in ['NAN','NONE','']: return None
            if not (c.startswith('IRAS') or c.startswith('IDS')): return None
            _, nome = classificar_unidade_sanitaria(c, loc, idade)
            return None if 'não identificada' in nome.lower() else nome

        if 'codigo_do_site' in df_graphs.columns:
            df_graphs['_nome_un'] = df_graphs.apply(
                lambda r: classif_dir(r.get('codigo_do_site'),
                                      r.get('local_de_colheita'),
                                      r.get('idade')), axis=1)
            df_val = df_graphs[df_graphs['_nome_un'].notna()].copy()
            rem = len(df_graphs) - len(df_val)
            if rem > 0:
                st.caption(f"ℹ️ {rem} amostra(s) sem código de sítio válido excluída(s) deste gráfico.")
            ug = df_val.groupby('_nome_un').agg(
                total=('codigo_do_site','count'),
                positivas=('_qualquer_pos','sum')).reset_index()
            ug['negativas'] = ug['total'] - ug['positivas']
            ug = ug.sort_values('total', ascending=True)
            if len(ug):
                n_units = len(ug)
                st.plotly_chart({
                    'data': [
                        {'y':ug['_nome_un'].tolist(),'x':ug['positivas'].tolist(),
                         'type':'bar','name':'Positivas','orientation':'h',
                         'marker':{'color':'#ef4444'},
                         'text':[str(v) if v > 0 else '' for v in ug['positivas'].tolist()],
                         'textposition':'inside','insidetextanchor':'middle'},
                        {'y':ug['_nome_un'].tolist(),'x':ug['negativas'].tolist(),
                         'type':'bar','name':'Negativas','orientation':'h',
                         'marker':{'color':'#22c55e'},
                         'text':[str(v) if v > 0 else '' for v in ug['negativas'].tolist()],
                         'textposition':'inside','insidetextanchor':'middle'},
                    ],
                    'layout': {
                        'title':{'text':'Distribuição de amostras por unidade sanitária','y':0.98},
                        'xaxis':{'title':'N.º de amostras'},
                        'yaxis':{'title':'','automargin':True},
                        'barmode':'stack',
                        'height':max(460, n_units*58 + 100),
                        'margin':{'l':320,'r':20,'t':80,'b':50},
                        'showlegend':True,
                        'legend':{
                            'orientation':'h',
                            'traceorder':'reversed',
                            'yanchor':'bottom','y':1.02,
                            'xanchor':'left','x':0,
                        }
                    }
                }, use_container_width=True, key='gc_unidades')
                with st.expander("📋 Tabela — detalhe por unidade"):
                    ug2 = ug[['_nome_un','total','positivas','negativas']].copy()
                    ug2['% positividade'] = (ug2['positivas']/ug2['total']*100).round(1)
                    ug2.columns = ['Unidade','Total','Positivas','Negativas','% Positividade']
                    st.dataframe(ug2, use_container_width=True, hide_index=True)

        # ── Gráfico 3: Tendência temporal ─────────────────────────────────
        st.markdown("---")
        st.header("2️⃣ Tendência Temporal — Positividade por Semana Epidemiológica")

        # Semanas presentes nos dados filtrados
        _df_sv = df_graphs[df_graphs['semana_epidemiologica'].notna()]
        semanas = sorted(
            _df_sv['semana_epidemiologica'].astype(str).unique(),
            key=lambda x: int(x) if x.isdigit() else 0)

        if not semanas:
            st.warning("Sem semanas epidemiológicas para o período seleccionado.")
        else:
            import plotly.graph_objects as go
            from plotly.subplots import make_subplots

            # Detectar semanas consecutivas — inserir None para quebrar a linha
            # quando existem semanas sem dados (ex: gap entre sem 6 e sem 49)
            def semanas_continuas(sems):
                """Expande lista de semanas para série contínua; None onde não há dados."""
                nums = [int(s) for s in sems]
                full = list(range(min(nums), max(nums)+1))
                presentes = set(nums)
                return [str(n) if n in presentes else None for n in full]

            sem_x = semanas_continuas(semanas) if len(semanas) > 1 else semanas

            # Painel Influenza + SARS
            st.subheader("Influenza e SARS-CoV-2")
            rows_sup = {}
            for sem in semanas:
                d = df_graphs[df_graphs['semana_epidemiologica'] == sem]
                flu_t = d['resultado_flu'].apply(lambda x: pd.notna(x) and str(x).strip().lower() not in ['','nan','none']).sum() if 'resultado_flu' in d.columns else 0
                sar_t = d['resultado_sars'].apply(lambda x: pd.notna(x) and str(x).strip().lower() not in ['','nan','none']).sum() if 'resultado_sars' in d.columns else 0
                flu_p = int(d['resultado_flu'].astype(str).str.upper().str.contains('POSITIVO', na=False).sum())  if 'resultado_flu'  in d.columns else 0
                sar_p = int(d['resultado_sars'].astype(str).str.upper().str.contains('POSITIVO', na=False).sum()) if 'resultado_sars' in d.columns else 0
                rows_sup[sem] = {'tot':len(d),
                                 'flu_pct': flu_p/flu_t*100 if flu_t else 0,
                                 'sar_pct': sar_p/sar_t*100 if sar_t else 0}

            # Construir séries com None nos gaps para quebrar a linha
            x_vals    = [s for s in sem_x]  # inclui None onde há gap
            tot_vals  = [rows_sup[s]['tot']     if s and s in rows_sup else 0    for s in sem_x]
            flu_vals  = [rows_sup[s]['flu_pct'] if s and s in rows_sup else None for s in sem_x]
            sar_vals  = [rows_sup[s]['sar_pct'] if s and s in rows_sup else None for s in sem_x]
            # Para os ticks, só mostrar as semanas com dados reais
            tick_vals = [s for s in sem_x if s is not None]

            fig_sup = make_subplots(specs=[[{"secondary_y": True}]])
            fig_sup.add_trace(go.Bar(x=x_vals, y=tot_vals, name='Amostras testadas',
                marker_color='#d4af37', opacity=0.55,
                hovertemplate='Semana %{x}<br>Testadas: %{y}<extra></extra>'), secondary_y=False)
            fig_sup.add_trace(go.Scatter(x=x_vals, y=flu_vals, name='% Influenza',
                mode='lines+markers', connectgaps=False,
                line=dict(color='#1e3a8a', width=3), marker=dict(size=8),
                hovertemplate='Sem. %{x}<br>Influenza: %{y:.1f}%<extra></extra>'), secondary_y=True)
            fig_sup.add_trace(go.Scatter(x=x_vals, y=sar_vals, name='% SARS-CoV-2',
                mode='lines+markers', connectgaps=False,
                line=dict(color='#dc2626', width=3), marker=dict(size=8),
                hovertemplate='Sem. %{x}<br>SARS-CoV-2: %{y:.1f}%<extra></extra>'), secondary_y=True)
            fig_sup.update_layout(
                title='Amostras testadas e positividade semanal — Influenza e SARS-CoV-2',
                xaxis=dict(
                    title='Semana epidemiológica', tickangle=-45,
                    tickmode='array', tickvals=tick_vals, ticktext=tick_vals,
                    categoryorder='array', categoryarray=tick_vals,
                ),
                height=440, hovermode='x unified',
                legend=dict(orientation='h', y=1.12, x=0),
                plot_bgcolor='white', paper_bgcolor='white')
            fig_sup.update_yaxes(title_text="N.º de amostras", secondary_y=False,
                showgrid=True, gridcolor='#e5e7eb',
                range=[0, max(max(tot_vals)*1.25, 10)])
            fig_sup.update_yaxes(title_text="% de positividade", secondary_y=True,
                showgrid=False,
                range=[0, max(max(v for v in flu_vals+sar_vals if v is not None)*1.25, 10)])
            st.plotly_chart(fig_sup, use_container_width=True, key='g_tempo_sup')

            # Painel RSV
            st.subheader("RSV")
            rows_rsv = {}
            for sem in semanas:
                d = df_graphs[df_graphs['semana_epidemiologica'] == sem]
                rsv_t = d['resultado_rsv'].apply(lambda x: pd.notna(x) and str(x).strip().lower() not in ['','nan','none']).sum() if 'resultado_rsv' in d.columns else 0
                rsv_p = int(d['resultado_rsv'].astype(str).str.upper().str.contains('POSITIVO', na=False).sum()) if 'resultado_rsv' in d.columns else 0
                rows_rsv[sem] = {'rsv_t': rsv_t, 'rsv_pct': rsv_p/rsv_t*100 if rsv_t else 0}

            rsv_t_vals   = [rows_rsv[s]['rsv_t']   if s and s in rows_rsv else 0    for s in sem_x]
            rsv_pct_vals = [rows_rsv[s]['rsv_pct'] if s and s in rows_rsv else None for s in sem_x]

            fig_rsv = make_subplots(specs=[[{"secondary_y": True}]])
            fig_rsv.add_trace(go.Bar(x=x_vals, y=rsv_t_vals, name='Amostras testadas (RSV)',
                marker_color='#d4af37', opacity=0.55,
                hovertemplate='Semana %{x}<br>RSV testadas: %{y}<extra></extra>'), secondary_y=False)
            fig_rsv.add_trace(go.Scatter(x=x_vals, y=rsv_pct_vals, name='% RSV',
                mode='lines+markers', connectgaps=False,
                line=dict(color='#dc2626', width=3), marker=dict(size=8),
                hovertemplate='Sem. %{x}<br>RSV: %{y:.1f}%<extra></extra>'), secondary_y=True)
            fig_rsv.update_layout(
                title='Amostras testadas e positividade semanal — RSV',
                xaxis=dict(
                    title='Semana epidemiológica', tickangle=-45,
                    tickmode='array', tickvals=tick_vals, ticktext=tick_vals,
                    categoryorder='array', categoryarray=tick_vals,
                ),
                height=400, hovermode='x unified',
                legend=dict(orientation='h', y=1.12, x=0),
                plot_bgcolor='white', paper_bgcolor='white')
            fig_rsv.update_yaxes(title_text="N.º de amostras (RSV)", secondary_y=False,
                showgrid=True, gridcolor='#e5e7eb',
                range=[0, max(max(rsv_t_vals)*1.25, 10)])
            fig_rsv.update_yaxes(title_text="% de positividade", secondary_y=True,
                showgrid=False,
                range=[0, max(max(v for v in rsv_pct_vals if v is not None)*1.25, 10)])
            st.plotly_chart(fig_rsv, use_container_width=True, key='g_tempo_rsv')

            # ── Gráfico 4: Subtipos Influenza ─────────────────────────────
            st.markdown("---")
            st.subheader("Circulação de Subtipos do Vírus Influenza")

            sub_data = {}
            for sem in semanas:
                d = df_graphs[df_graphs['semana_epidemiologica'] == sem]
                cnt = {'A(H3N2)':0,'A(H1pdm)':0,'B(Victoria)':0,'B(Yamagata)':0}
                for _, row in d.iterrows():
                    if 'POSITIVO' not in str(row.get('resultado_flu','')).upper():
                        continue
                    for col_s, lbl in [('h3','A(H3N2)'),('apdm','A(H1pdm)'),
                                       ('h1pdm','A(H1pdm)'),('vic','B(Victoria)'),('yam','B(Yamagata)')]:
                        val = row.get(col_s)
                        if val is not None and pd.notna(val):
                            try:
                                if float(val) < 40: cnt[lbl] += 1
                            except (ValueError, TypeError):
                                if 'POSITIVO' in str(val).upper(): cnt[lbl] += 1
                sub_data[sem] = cnt

            cores_sub = {'A(H3N2)':'#7c3aed','A(H1pdm)':'#3b82f6',
                         'B(Victoria)':'#ef4444','B(Yamagata)':'#f97316'}
            fig_sub   = go.Figure()
            has_sub   = False
            for sub, cor in cores_sub.items():
                y_vals = [sub_data[s][sub] if s and s in sub_data else None for s in sem_x]
                if any(v is not None and v > 0 for v in y_vals):
                    has_sub = True
                    fig_sub.add_trace(go.Scatter(
                        x=x_vals, y=y_vals, name=sub,
                        mode='lines+markers', connectgaps=False,
                        line=dict(color=cor, width=2), marker=dict(size=8),
                        hovertemplate=f'{sub} — Sem. %{{x}}: %{{y}} casos<extra></extra>'))
            if has_sub:
                fig_sub.update_layout(
                    title='Circulação semanal de subtipos do vírus Influenza — Moçambique',
                    xaxis=dict(
                        title='Semana epidemiológica', tickangle=-45,
                        tickmode='array', tickvals=tick_vals, ticktext=tick_vals,
                        categoryorder='array', categoryarray=tick_vals,
                    ),
                    yaxis=dict(title='Casos positivos', showgrid=True, gridcolor='#e5e7eb'),
                    height=400, hovermode='x unified',
                    plot_bgcolor='white', paper_bgcolor='white',
                    legend=dict(orientation='h', y=1.08, x=0))
                st.plotly_chart(fig_sub, use_container_width=True, key='g_subtipos')
            else:
                st.info("ℹ️ Sem dados de subtipos de Influenza para o período seleccionado.")

        # ── Gráfico 5: Distribuição etária ────────────────────────────────
        st.markdown("---")
        st.header("3️⃣ Distribuição por Faixa Etária")

        if 'idade' in df_graphs.columns:
            df_graphs['_idade_a'] = df_graphs['idade'].apply(extrair_valor_idade)
            df_graphs['_faixa']   = pd.cut(
                df_graphs['_idade_a'],
                bins=[0,2,5,15,50,65,150],
                labels=['0–<2','2–<5','5–<15','15–<50','50–<65','≥65'],
                include_lowest=True)

            df_et = df_graphs[df_graphs['_qualquer_pos']].copy()
            if len(df_et) > 0:
                et_rows = []
                for _, r in df_et.iterrows():
                    f = r.get('_faixa')
                    if pd.isna(f): continue
                    if 'resultado_flu' in r and 'POSITIVO' in str(r['resultado_flu']).upper():
                        if 'infa' in r and 'POSITIVO' in str(r['infa']).upper():
                            et_rows.append({'faixa':str(f),'virus':'Influenza A'})
                        if 'infb' in r and 'POSITIVO' in str(r['infb']).upper():
                            et_rows.append({'faixa':str(f),'virus':'Influenza B'})
                    if 'resultado_sars' in r and 'POSITIVO' in str(r.get('resultado_sars','')).upper():
                        et_rows.append({'faixa':str(f),'virus':'SARS-CoV-2'})
                    if 'resultado_rsv' in r and 'POSITIVO' in str(r.get('resultado_rsv','')).upper():
                        et_rows.append({'faixa':str(f),'virus':'RSV'})
                if et_rows:
                    df_fe    = pd.DataFrame(et_rows)
                    cnt_fe   = df_fe.groupby(['faixa','virus']).size().reset_index(name='n')
                    cores_et = {'Influenza A':'#3b82f6','Influenza B':'#818cf8',
                                'SARS-CoV-2':'#ef4444','RSV':'#f59e0b'}
                    fig_et = {'data':[],'layout':{
                        'title':'Casos positivos por faixa etária e patógeno',
                        'xaxis':{'title':'Grupo etário'},
                        'yaxis':{'title':'Casos positivos'},
                        'barmode':'group','height':460}}
                    for vir in cnt_fe['virus'].unique():
                        dv = cnt_fe[cnt_fe['virus']==vir]
                        fig_et['data'].append({'x':dv['faixa'].tolist(),'y':dv['n'].tolist(),
                            'type':'bar','name':vir,
                            'marker':{'color':cores_et.get(vir,'#6b7280')}})
                    st.plotly_chart(fig_et, use_container_width=True, key='g_etaria')
                    st.subheader("📋 Resumo por faixa etária")
                    tbl_et = df_fe.groupby(['faixa','virus']).size().unstack(fill_value=0)
                    st.dataframe(tbl_et, use_container_width=True)
                else:
                    st.warning("Sem dados etários disponíveis para os casos positivos.")
            else:
                st.warning("Sem casos positivos no período/filtro seleccionado.")

# ============================================================================
# RODAPÉ
# ============================================================================

st.markdown("---")
st.caption(
    f"**Sistema de Vigilância de IRAs — Instituto Nacional de Saúde, Moçambique**  "
    f"v3.0 · Influenza · RSV · SARS-CoV-2 · "
    f"Actualizado: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
)
